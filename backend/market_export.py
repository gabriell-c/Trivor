"""Market intelligence export functions - MD, PDF, DOCX"""

from io import BytesIO
import json
from datetime import datetime


def generate_market_markdown_export(report: dict, job_title: str, seniority: str, location: str, model_name: str) -> str:
    """Generate a well-structured Markdown export of the market analysis report."""
    s = report.get("summary", {})
    stats = report.get("statistics", {})
    jobs = report.get("sample_jobs", [])
    now_str = s.get("generated_at", datetime.now().strftime("%Y-%m-%d %H:%M"))

    lines = []
    # Header
    lines.append(f"# Intelligence de Mercado — {job_title}")
    lines.append(f"")
    lines.append(f"**Seniority:** {seniority}  |  **Stack:** {', '.join(s.get('target_stack', []))}  |  **Local:** {location}")
    lines.append(f"**Gerado em:** {now_str}  |  **Modelo:** {model_name}")
    lines.append(f"")

    # Summary metrics
    lines.append(f"## Resumo")
    lines.append(f"")
    lines.append(f"| Métrica | Valor |")
    lines.append(f"|---|---|")
    lines.append(f"| Vagas coletadas | {s.get('total_jobs_scanned', 0)} |")
    lines.append(f"| Pré-filtradas | {s.get('pre_filtered_count', 0)} |")
    lines.append(f"| Analisadas com IA | {s.get('relevant_jobs_analyzed', 0)} |")
    lines.append(f"| Descartadas | {s.get('discarded_jobs', 0)} |")
    lines.append(f"| Confiança | {s.get('confidence_score', 'N/A')} — {s.get('confidence_reason', '')} |")
    lines.append(f"| Anos de experiência (mediana) | {stats.get('exp_years_median', 'N/A')} |")
    lines.append(f"")

    # Technologies
    if stats.get("required_technologies"):
        lines.append(f"## Tecnologias Obrigatórias")
        lines.append(f"")
        lines.append(f"| Tecnologia | Ocorrências | % |")
        lines.append(f"|---|---|---|")
        for t in stats["required_technologies"]:
            lines.append(f"| {t['name']} | {t['count']} | {t['percentage']}% |")
        lines.append(f"")

    if stats.get("desirable_technologies"):
        lines.append(f"## Tecnologias Desejáveis / Diferenciais")
        lines.append(f"")
        lines.append(f"| Tecnologia | Ocorrências | % |")
        lines.append(f"|---|---|---|")
        for t in stats["desirable_technologies"]:
            lines.append(f"| {t['name']} | {t['count']} | {t['percentage']}% |")
        lines.append(f"")

    # Experience distribution
    if stats.get("exp_years_distribution"):
        lines.append(f"## Distribuição de Anos de Experiência")
        lines.append(f"")
        lines.append(f"| Faixa | Vagas |")
        lines.append(f"|---|---|")
        for range_key, count in sorted(stats["exp_years_distribution"].items()):
            lines.append(f"| {range_key} | {count} |")
        lines.append(f"")

    # Modalities
    if stats.get("modalities"):
        lines.append(f"## Modalidades")
        lines.append(f"")
        lines.append(f"| Modalidade | Vagas | % |")
        lines.append(f"|---|---|---|")
        for m in stats["modalities"]:
            lines.append(f"| {m['name']} | {m['count']} | {m['percentage']}% |")
        lines.append(f"")

    # Soft skills
    if stats.get("top_soft_skills"):
        lines.append(f"## Soft Skills mais pedidas")
        lines.append(f"")
        lines.append(f"| Skill | Ocorrências |")
        lines.append(f"|---|---|")
        for sk in stats["top_soft_skills"][:15]:
            lines.append(f"| {sk['name']} | {sk['count']} |")
        lines.append(f"")

    # Certifications
    if stats.get("top_certifications"):
        lines.append(f"## Certificações mais pedidas")
        lines.append(f"")
        lines.append(f"| Certificação | Ocorrências |")
        lines.append(f"|---|---|")
        for c in stats["top_certifications"][:10]:
            lines.append(f"| {c['name']} | {c['count']} |")
        lines.append(f"")

    # Sample jobs
    lines.append(f"## Vagas Analisadas ({len(jobs)} relevantes)")
    lines.append(f"")
    for i, job in enumerate(jobs, 1):
        lines.append(f"### {i}. {job.get('title', 'N/A')} — {job.get('company', 'N/A')}")
        lines.append(f"")
        lines.append(f"- **Local:** {job.get('location', 'N/A')}  |  **Modalidade:** {job.get('modality', 'N/A')}  |  **Fonte:** {job.get('source', 'N/A')}")
        reqs = job.get('requirements', [])
        if reqs:
            lines.append(f"- **Req. técnicas:** {', '.join(reqs)}")
        nice = job.get('nice_to_have', [])
        if nice:
            lines.append(f"- **Nice-to-have:** {', '.join(nice)}")
        soft = job.get('soft_skills', [])
        if soft:
            lines.append(f"- **Soft skills:** {', '.join(soft)}")
        certs = job.get('certifications', [])
        if certs:
            lines.append(f"- **Certificações:** {', '.join(certs)}")
        salary_min = job.get('salary_min')
        salary_max = job.get('salary_max')
        currency = job.get('currency')
        if salary_min is not None and salary_max is not None:
            lines.append(f"- **Salário:** {currency or 'R$'} {salary_min} — {salary_max}")
        level = job.get('role_level')
        if level:
            lines.append(f"- **Nível:** {level}")
        lines.append(f"")

    return "\n".join(lines)


def generate_market_docx_export(report: dict, job_title: str, seniority: str, location: str, model_name: str) -> BytesIO:
    """Generate a well-formatted DOCX export of the market analysis report."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    s = report.get("summary", {})
    stats = report.get("statistics", {})
    jobs = report.get("sample_jobs", [])
    now_str = s.get("generated_at", datetime.now().strftime("%d/%m/%Y %H:%M"))

    # ── Header ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("TRIVOR")
    r.font.name = "Arial"
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p = doc.add_paragraph()
    r = p.add_run(f"Inteligência de Mercado — {job_title}")
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x63, 0x66, 0x1)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Meta info line
    p = doc.add_paragraph()
    r = p.add_run(f"Seniority: {seniority}  |  Stack: {', '.join(s.get('target_stack', []))}  |  Local: {location}  |  Modelo: {model_name}")
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    p = doc.add_paragraph()
    r = p.add_run(f"Gerado em: {now_str}")
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ── Summary table ──
    p = doc.add_paragraph()
    r = p.add_run("Resumo")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    summary_data = [
        ("Vagas coletadas", str(s.get("total_jobs_scanned", 0))),
        ("Pré-filtradas", str(s.get("pre_filtered_count", 0))),
        ("Analisadas com IA", str(s.get("relevant_jobs_analyzed", 0))),
        ("Descartadas", str(s.get("discarded_jobs", 0))),
        ("Confiança", f"{s.get('confidence_score', 'N/A')} — {s.get('confidence_reason', '')}"),
        ("Exp. média (mediana)", str(stats.get("exp_years_median", "N/A"))),
    ]
    t = doc.add_table(rows=len(summary_data), cols=2)
    t.style = "Table Grid"
    for i, (label, value) in enumerate(summary_data):
        t.rows[i].cells[0].text = label
        t.rows[i].cells[1].text = value
        for cell in t.rows[i].cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
            cell.paragraphs[0].runs[0].font.bold = (cell == t.rows[i].cells[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ── Helper for section headers ──
    def add_section(title: str):
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.font.name = "Arial"
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_subsection(title: str):
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def add_bullet(text: str):
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(10)

    def make_table(headers, rows_data):
        if not rows_data:
            return
        t = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
        t.style = "Table Grid"
        # Header row
        for j, h in enumerate(headers):
            cell = t.rows[0].cells[j]
            cell.text = h
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xF8, 0xFA, 0xFC)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "0F172A")
            cell._element.get_or_add_tcPr().append(shd)
        # Data rows
        for i, row in enumerate(rows_data):
            for j, val in enumerate(row):
                cell = t.rows[i + 1].cells[j]
                cell.text = str(val)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in para.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(9)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Technologies ──
    if stats.get("required_technologies"):
        add_section("Tecnologias Obrigatórias")
        rows = [(t["name"], str(t["count"]), f"{t['percentage']}%") for t in stats["required_technologies"][:20]]
        make_table(["Tecnologia", "Ocorrências", "%"], rows)

    if stats.get("desirable_technologies"):
        add_section("Tecnologias Desejáveis")
        rows = [(t["name"], str(t["count"]), f"{t['percentage']}%") for t in stats["desirable_technologies"][:15]]
        make_table(["Tecnologia", "Ocorrências", "%"], rows)

    # ── Experience distribution ──
    if stats.get("exp_years_distribution"):
        add_section("Distribuição de Anos de Experiência")
        rows = [(k, str(v)) for k, v in sorted(stats["exp_years_distribution"].items())]
        make_table(["Faixa", "Vagas"], rows)

    # ── Modalities ──
    if stats.get("modalities"):
        add_section("Modalidades")
        rows = [(m["name"], str(m["count"]), f"{m['percentage']}%") for m in stats["modalities"]]
        make_table(["Modalidade", "Vagas", "%"], rows)

    # ── Soft skills ──
    if stats.get("top_soft_skills"):
        add_section("Soft Skills mais pedidas")
        rows = [(sk["name"], str(sk["count"])) for sk in stats["top_soft_skills"][:15]]
        make_table(["Skill", "Ocorrências"], rows)

    # ── Certifications ──
    if stats.get("top_certifications"):
        add_section("Certificações mais pedidas")
        rows = [(c["name"], str(c["count"])) for c in stats["top_certifications"][:10]]
        make_table(["Certificação", "Ocorrências"], rows)

    # ── Sample jobs ──
    if jobs:
        add_section(f"Vagas Analisadas ({len(jobs)} relevantes)")

        for i, job in enumerate(jobs, 1):
            p = doc.add_paragraph()
            r = p.add_run(f"{i}. {job.get('title', 'N/A')} — {job.get('company', 'N/A')}")
            r.font.name = "Arial"
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

            info = f"Local: {job.get('location', 'N/A')}  |  Modalidade: {job.get('modality', 'N/A')}  |  Fonte: {job.get('source', 'N/A')}"
            p = doc.add_paragraph()
            r = p.add_run(info)
            r.font.name = "Arial"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

            reqs = job.get("requirements", [])
            if reqs:
                p = doc.add_paragraph()
                r = p.add_run("Req. técnicas: ")
                r.font.name = "Arial"
                r.font.size = Pt(9)
                r.font.bold = True
                r = p.add_run(", ".join(reqs))
                r.font.name = "Arial"
                r.font.size = Pt(9)

            nice = job.get("nice_to_have", [])
            if nice:
                p = doc.add_paragraph()
                r = p.add_run("Nice-to-have: ")
                r.font.name = "Arial"
                r.font.size = Pt(9)
                r.font.bold = True
                r = p.add_run(", ".join(nice))
                r.font.name = "Arial"
                r.font.size = Pt(9)

            soft = job.get("soft_skills", [])
            if soft:
                p = doc.add_paragraph()
                r = p.add_run("Soft skills: ")
                r.font.name = "Arial"
                r.font.size = Pt(9)
                r.font.bold = True
                r = p.add_run(", ".join(soft))
                r.font.name = "Arial"
                r.font.size = Pt(9)

            certs = job.get("certifications", [])
            if certs:
                p = doc.add_paragraph()
                r = p.add_run("Certificações: ")
                r.font.name = "Arial"
                r.font.size = Pt(9)
                r.font.bold = True
                r = p.add_run(", ".join(certs))
                r.font.name = "Arial"
                r.font.size = Pt(9)

            salary_min = job.get("salary_min")
            salary_max = job.get("salary_max")
            currency = job.get("currency")
            if salary_min is not None and salary_max is not None:
                p = doc.add_paragraph()
                r = p.add_run(f"Salário: {currency or 'R$'} {salary_min} — {salary_max}")
                r.font.name = "Arial"
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x05, 0x96, 0x69)

            level = job.get("role_level")
            if level:
                p = doc.add_paragraph()
                r = p.add_run(f"Nível: {level}")
                r.font.name = "Arial"
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x63, 0x66, 0x1)

            # Add a small spacer between jobs
            p = doc.add_paragraph()
            r = p.add_run("")
            doc.add_paragraph().paragraph_format.space_after = Pt(8)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_market_pdf_export(report: dict, job_title: str, seniority: str, location: str, model_name: str) -> BytesIO:
    """Generate a modern, well-formatted PDF export of the market analysis report."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, KeepTogether
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.8*cm,
        leftMargin=1.8*cm,
        topMargin=1.5*cm,
        bottomMargin=1.5*cm,
    )

    styles = getSampleStyleSheet()

    # Try to register a Brazilian font for proper accent characters
    font_paths = [
        os.path.join(os.environ.get('WINDIR', 'C:\\Windows'), 'Fonts', 'arial.ttf'),
        os.path.join(os.environ.get('WINDIR', 'C:\\Windows'), 'Fonts', 'ARIAL.TTF'),
    ]
    font_registered = False
    for fp in font_paths:
        if os.path.exists(fp):
            try:
                pdfmetrics.registerFont(TTFont('Arial', fp))
                font_registered = True
                break
            except Exception:
                pass

    # Also try to register bold variant
    bold_registered = False
    bold_paths = [
        os.path.join(os.environ.get('WINDIR', 'C:\\Windows'), 'Fonts', 'arialbd.ttf'),
        os.path.join(os.environ.get('WINDIR', 'C:\\Windows'), 'Fonts', 'ARIALBD.TTF'),
    ]
    for bp in bold_paths:
        if os.path.exists(bp):
            try:
                pdfmetrics.registerFont(TTFont('Arial-Bold', bp))
                bold_registered = True
                break
            except Exception:
                pass

    font_name = 'Arial' if font_registered else 'Helvetica'
    bold_font = 'Arial-Bold' if bold_registered else 'Helvetica-Bold'

    # Custom styles
    title_style = ParagraphStyle(
        'MktTitle',
        parent=styles['Normal'],
        fontName=bold_font,
        fontSize=18,
        textColor=colors.HexColor('#0F172A'),
        spaceAfter=2,
        leading=22,
    )

    subtitle_style = ParagraphStyle(
        'MktSubtitle',
        parent=styles['Normal'],
        fontName=bold_font,
        fontSize=8,
        textColor=colors.HexColor('#6366F1'),
        spaceAfter=1,
        leading=10,
    )

    meta_style = ParagraphStyle(
        'MktMeta',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=8,
        textColor=colors.HexColor('#64748B'),
        spaceAfter=1,
        leading=10,
    )

    section_style = ParagraphStyle(
        'MktSection',
        parent=styles['Normal'],
        fontName=bold_font,
        fontSize=11,
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=10,
        spaceAfter=4,
        leading=14,
    )

    subsection_style = ParagraphStyle(
        'MktSubsection',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=9,
        textColor=colors.HexColor('#475569'),
        spaceBefore=6,
        spaceAfter=2,
        leading=12,
    )

    body_style = ParagraphStyle(
        'MktBody',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=8.5,
        textColor=colors.HexColor('#334155'),
        spaceAfter=2,
        leading=11,
    )

    small_style = ParagraphStyle(
        'MktSmall',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=7.5,
        textColor=colors.HexColor('#64748B'),
        spaceAfter=1,
        leading=10,
    )

    # Accent table style helper
    def make_table_style(header_fill="0F172A", alt_fill="F8FAFC"):
        return TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 7.5),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#334155')),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E2E8F0')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor(f'#{alt_fill}'), colors.HexColor('#FFFFFF')]),
            ('TOPPADDING', (0, 0), (-1, -1), 3),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
            ('LEFTPADDING', (0, 0), (-1, -1), 5),
            ('RIGHTPADDING', (0, 0), (-1, -1), 5),
        ])

    def header_style(header_fill="0F172A"):
        return TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), bold_font),
            ('FONTSIZE', (0, 0), (-1, -1), 7.5),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#F8FAFC')),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(f'#{header_fill}')),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('LEFTPADDING', (0, 0), (-1, -1), 5),
            ('RIGHTPADDING', (0, 0), (-1, -1), 5),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#1E293B')),
        ])

    elements = []

    s = report.get("summary", {})
    stats = report.get("statistics", {})
    jobs = report.get("sample_jobs", [])
    now_str = s.get("generated_at", datetime.now().strftime("%d/%m/%Y %H:%M"))

    # ── Header ──
    elements.append(Paragraph("TRIVOR", title_style))
    elements.append(Paragraph(f"Inteligência de Mercado — {job_title}", subtitle_style))
    meta_text = f"Seniority: {seniority}  |  Stack: {', '.join(s.get('target_stack', []))}  |  Local: {location}"
    elements.append(Paragraph(meta_text, meta_style))
    elements.append(Paragraph(f"Modelo: {model_name}  |  Gerado em: {now_str}", meta_style))
    elements.append(Spacer(1, 0.3*cm))

    # ── Summary table ──
    elements.append(Paragraph("Resumo", section_style))
    summary_rows = [
        ["Vagas coletadas", str(s.get("total_jobs_scanned", 0))],
        ["Pré-filtradas", str(s.get("pre_filtered_count", 0))],
        ["Analisadas com IA", str(s.get("relevant_jobs_analyzed", 0))],
        ["Descartadas", str(s.get("discarded_jobs", 0))],
        ["Confiança", f"{s.get('confidence_score', 'N/A')} — {s.get('confidence_reason', '')}"],
        ["Exp. média (mediana)", str(stats.get("exp_years_median", "N/A"))],
    ]
    t = Table(summary_rows, colWidths=[4.5*cm, 12*cm])
    t.setStyle(header_style("1E293B"))
    t.setStyle(make_table_style("1E293B", "F1F5F9"))
    elements.append(t)
    elements.append(Spacer(1, 0.3*cm))

    # ── Technologies ──
    if stats.get("required_technologies"):
        elements.append(Paragraph("Tecnologias Obrigatórias", section_style))
        header = [Paragraph(h, ParagraphStyle('h', fontName=bold_font, fontSize=7.5, textColor=colors.white)) for h in ["Tecnologia", "Ocorrências", "%"]]
        data = [[Paragraph(t["name"], body_style),
                 Paragraph(str(t["count"]), ParagraphStyle('n', alignment=1)),
                 Paragraph(f"{t['percentage']}%", ParagraphStyle('p', alignment=1))]
                for t in stats["required_technologies"][:20]]
        rows = [header] + data
        t = Table(rows, colWidths=[6*cm, 2.5*cm, 2*cm])
        t.setStyle(header_style("4F46E5"))
        t.setStyle(make_table_style("4F46E5", "F8FAFC"))
        elements.append(t)
        elements.append(Spacer(1, 0.2*cm))

    if stats.get("desirable_technologies"):
        elements.append(Paragraph("Tecnologias Desejáveis", section_style))
        header = [Paragraph(h, ParagraphStyle('h', fontName=bold_font, fontSize=7.5, textColor=colors.white)) for h in ["Tecnologia", "Ocorrências", "%"]]
        data = [[Paragraph(t["name"], body_style),
                 Paragraph(str(t["count"]), ParagraphStyle('n', alignment=1)),
                 Paragraph(f"{t['percentage']}%", ParagraphStyle('p', alignment=1))]
                for t in stats["desirable_technologies"][:15]]
        rows = [header] + data
        t = Table(rows, colWidths=[6*cm, 2.5*cm, 2*cm])
        t.setStyle(header_style("0891B2"))
        t.setStyle(make_table_style("0891B2", "F0FDFA"))
        elements.append(t)
        elements.append(Spacer(1, 0.2*cm))

    # ── Experience distribution ──
    if stats.get("exp_years_distribution"):
        elements.append(Paragraph("Distribuição de Anos de Experiência", section_style))
        header = [Paragraph(h, ParagraphStyle('h', fontName=bold_font, fontSize=7.5, textColor=colors.white)) for h in ["Faixa", "Vagas"]]
        data = [[Paragraph(k, body_style), Paragraph(str(v), ParagraphStyle('n', alignment=1))]
                for k, v in sorted(stats["exp_years_distribution"].items())]
        rows = [header] + data
        t = Table(rows, colWidths=[6*cm, 5*cm])
        t.setStyle(header_style("059669"))
        t.setStyle(make_table_style("059669", "F0FDF4"))
        elements.append(t)
        elements.append(Spacer(1, 0.2*cm))

    # ── Modalities ──
    if stats.get("modalities"):
        elements.append(Paragraph("Modalidades", section_style))
        header = [Paragraph(h, ParagraphStyle('h', fontName=bold_font, fontSize=7.5, textColor=colors.white)) for h in ["Modalidade", "Vagas", "%"]]
        data = [[Paragraph(m["name"], body_style),
                 Paragraph(str(m["count"]), ParagraphStyle('n', alignment=1)),
                 Paragraph(f"{m['percentage']}%", ParagraphStyle('p', alignment=1))]
                for m in stats["modalities"]]
        rows = [header] + data
        t = Table(rows, colWidths=[5*cm, 2.5*cm, 2*cm])
        t.setStyle(header_style("7C3AED"))
        t.setStyle(make_table_style("7C3AED", "FAF5FF"))
        elements.append(t)
        elements.append(Spacer(1, 0.2*cm))

    # ── Soft skills ──
    if stats.get("top_soft_skills"):
        elements.append(Paragraph("Soft Skills mais pedidas", section_style))
        header = [Paragraph(h, ParagraphStyle('h', fontName=bold_font, fontSize=7.5, textColor=colors.white)) for h in ["Skill", "Ocorrências"]]
        data = [[Paragraph(sk["name"], body_style),
                 Paragraph(str(sk["count"]), ParagraphStyle('n', alignment=1))]
                for sk in stats["top_soft_skills"][:15]]
        rows = [header] + data
        t = Table(rows, colWidths=[7*cm, 3*cm])
        t.setStyle(header_style("DC2626"))
        t.setStyle(make_table_style("DC2626", "FEF2F2"))
        elements.append(t)
        elements.append(Spacer(1, 0.2*cm))

    # ── Certifications ──
    if stats.get("top_certifications"):
        elements.append(Paragraph("Certificações mais pedidas", section_style))
        header = [Paragraph(h, ParagraphStyle('h', fontName=bold_font, fontSize=7.5, textColor=colors.white)) for h in ["Certificação", "Ocorrências"]]
        data = [[Paragraph(c["name"], body_style),
                 Paragraph(str(c["count"]), ParagraphStyle('n', alignment=1))]
                for c in stats["top_certifications"][:10]]
        rows = [header] + data
        t = Table(rows, colWidths=[7*cm, 3*cm])
        t.setStyle(header_style("D97706"))
        t.setStyle(make_table_style("D97706", "FFFBEB"))
        elements.append(t)
        elements.append(Spacer(1, 0.2*cm))

    # ── Sample jobs ──
    if jobs:
        elements.append(Paragraph(f"Vagas Analisadas ({len(jobs)} relevantes)", section_style))

        for i, job in enumerate(jobs, 1):
            # Keep each job together
            job_elements = []

            # Job title
            job_title_style = ParagraphStyle(
                'JobTitle',
                parent=styles['Normal'],
                fontName=bold_font,
                fontSize=9.5,
                textColor=colors.HexColor('#0F172A'),
                spaceBefore=6,
                spaceAfter=2,
            )
            job_elements.append(Paragraph(
                f"{i}. {job.get('title', 'N/A')} — {job.get('company', 'N/A')}",
                job_title_style
            ))

            # Meta info
            info_text = f"Local: {job.get('location', 'N/A')}  |  Modalidade: {job.get('modality', 'N/A')}  |  Fonte: {job.get('source', 'N/A')}"
            job_elements.append(Paragraph(info_text, meta_style))

            # Requirements
            reqs = job.get("requirements", [])
            if reqs:
                p = Paragraph("Req. técnicas: ", body_style)
                p2 = Paragraph(", ".join(reqs), body_style)
                job_elements.append(Paragraph("<b>Req. técnicas:</b> " + ", ".join(reqs), body_style))

            nice = job.get("nice_to_have", [])
            if nice:
                job_elements.append(Paragraph("<b>Nice-to-have:</b> " + ", ".join(nice), body_style))

            soft = job.get("soft_skills", [])
            if soft:
                job_elements.append(Paragraph("<b>Soft skills:</b> " + ", ".join(soft), body_style))

            certs = job.get("certifications", [])
            if certs:
                job_elements.append(Paragraph("<b>Certificações:</b> " + ", ".join(certs), body_style))

            salary_min = job.get("salary_min")
            salary_max = job.get("salary_max")
            currency = job.get("currency")
            if salary_min is not None and salary_max is not None:
                sal_style = ParagraphStyle(
                    'Salary',
                    parent=body_style,
                    textColor=colors.HexColor('#059669'),
                )
                job_elements.append(Paragraph(f"<b>Salário:</b> {currency or 'R$'} {salary_min} — {salary_max}", sal_style))

            level = job.get("role_level")
            if level:
                job_elements.append(Paragraph(f"<b>Nível:</b> {level}", body_style))

            # Spacer after each job
            job_elements.append(Spacer(1, 3))

            elements.append(KeepTogether(job_elements))

    doc.build(elements)
    buffer.seek(0)
    return buffer
