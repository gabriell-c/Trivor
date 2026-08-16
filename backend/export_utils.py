from io import BytesIO
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def generate_markdown_export(data: dict, filename: str, job_target: str, model_name: str) -> str:
    md = []
    md.append(f"# Trivor")
    md.append(f"**Arquivo:** {filename} | **Vaga / Nível:** {job_target} | **Modelo:** {model_name}\n")
    md.append(f"## 🏆 Resume Score: {data.get('nota', 0)} / 10\n")
    md.append(f"### 📋 Parecer Executivo\n{data.get('resumo_executivo', 'Sem resumo disponível.')}\n")
    
    md.append("## 🔍 Diagnóstico Estruturado por Seção\n")
    secoes = data.get('diagnostico_por_secao', {})
    titulos = {
        'dados_pessoais': '👤 Dados Pessoais & Contatos',
        'resumo_profissional': '📝 Resumo / Perfil Profissional',
        'experiencia_profissional': '💼 Experiência Profissional & Métricas',
        'educacao_e_cursos': '🎓 Formação Acadêmica & Cursos',
        'habilidades_e_keywords': '🛠️ Habilidades & Palavras-Chave'
    }
    for k, title in titulos.items():
        if k in secoes and secoes[k]:
            item = secoes[k]
            status = item.get('status', 'atencao').upper()
            md.append(f"### {title} [{status}]")
            if item.get('problema'):
                md.append(f"- **Problema Detectado:** {item['problema']}")
            if item.get('como_corrigir'):
                md.append(f"- **Como Ajustar:** {item['como_corrigir']}")
            md.append("")

    ats = data.get('analise_ats', {})
    if isinstance(ats, dict):
        md.append("## 🤖 Compatibilidade ATS (Robôs de Triagem)\n")
        md.append(f"- **ATS Score:** {ats.get('score_ats', 'N/A')}/10")
        md.append(f"- **Veredito:** {ats.get('veredito_robos', '')}\n")
        if ats.get('palavras_chave_faltantes'):
            md.append(f"**Palavras-chave Faltantes:** {', '.join(ats['palavras_chave_faltantes'])}\n")
        if ats.get('gargalos_formatacao'):
            md.append("**Gargalos de Formatação:**")
            for g in ats['gargalos_formatacao']:
                md.append(f"- {g}")
            md.append("")

    fortes = data.get('pontos_fortes', [])
    if fortes:
        md.append("## 📈 Pontos Fortes Rastreados")
        for f in fortes:
            md.append(f"- {f}")
        md.append("")

    tokens = data.get('uso_tokens')
    if tokens:
        md.append(f"---\n*Métricas de Consumo de Tokens: Prompt: {tokens.get('prompt_tokens')} | Resposta: {tokens.get('completion_tokens')} | Total: {tokens.get('total_tokens')}*")

    return "\n".join(md)


def generate_docx_export(data: dict, filename: str, job_target: str, model_name: str) -> BytesIO:
    doc = Document()

    # Definir margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Estilo de Título Principal
    p_title = doc.add_paragraph()
    run_title = p_title.add_run("TRIVOR")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42) # Slate-900

    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run(f"Trivor  •  {filename}  •  {job_target}")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(10)
    run_sub.font.color.rgb = RGBColor(99, 102, 241) # Indigo-500

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Tabela de Score
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cell_score, cell_text = table.rows[0].cells
    cell_score.width = Inches(2.2)
    cell_text.width = Inches(4.6)

    set_cell_background(cell_score, "0F172A") # Slate 900
    set_cell_background(cell_text, "1E293B") # Slate 800

    p_sc = cell_score.paragraphs[0]
    p_sc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sc1 = p_sc.add_run(f"\n{data.get('nota', 0.0):.1f}\n")
    r_sc1.font.name = "Arial"
    r_sc1.font.size = Pt(36)
    r_sc1.font.bold = True
    r_sc1.font.color.rgb = RGBColor(99, 102, 241)

    r_sc2 = p_sc.add_run("RESUME SCORE / 10\n")
    r_sc2.font.name = "Arial"
    r_sc2.font.size = Pt(9)
    r_sc2.font.bold = True
    r_sc2.font.color.rgb = RGBColor(148, 163, 184)

    p_tx = cell_text.paragraphs[0]
    r_tx_h = p_tx.add_run("PARECER EXECUTIVO DO DIAGNÓSTICO\n\n")
    r_tx_h.font.name = "Arial"
    r_tx_h.font.size = Pt(10)
    r_tx_h.font.bold = True
    r_tx_h.font.color.rgb = RGBColor(226, 232, 240)

    r_tx_b = p_tx.add_run(data.get('resumo_executivo', ''))
    r_tx_b.font.name = "Arial"
    r_tx_b.font.size = Pt(10)
    r_tx_b.font.color.rgb = RGBColor(203, 213, 225)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # Diagnóstico por Seção
    h_sec = doc.add_paragraph()
    r_hsec = h_sec.add_run("DIAGNÓSTICO ESTRUTURADO POR SEÇÃO DO CURRÍCULO")
    r_hsec.font.name = "Arial"
    r_hsec.font.size = Pt(12)
    r_hsec.font.bold = True
    r_hsec.font.color.rgb = RGBColor(15, 23, 42)

    secoes = data.get('diagnostico_por_secao', {})
    titulos = {
        'dados_pessoais': 'Dados Pessoais & Contatos',
        'resumo_profissional': 'Resumo / Perfil Profissional',
        'experiencia_profissional': 'Experiência Profissional & Métricas',
        'educacao_e_cursos': 'Formação Acadêmica & Cursos',
        'habilidades_e_keywords': 'Habilidades & Palavras-Chave'
    }

    for k, title in titulos.items():
        if k in secoes and secoes[k]:
            item = secoes[k]
            p_s = doc.add_paragraph()
            r_st = p_s.add_run(f"• {title}  [{item.get('status', 'atencao').upper()}]\n")
            r_st.font.name = "Arial"
            r_st.font.size = Pt(11)
            r_st.font.bold = True
            r_st.font.color.rgb = RGBColor(30, 41, 59)

            if item.get('problema'):
                p_prob = doc.add_paragraph()
                p_prob.paragraph_format.left_indent = Inches(0.2)
                r_pr = p_prob.add_run(f"Problema Detectado: {item['problema']}")
                r_pr.font.name = "Arial"
                r_pr.font.size = Pt(9.5)
                r_pr.font.color.rgb = RGBColor(225, 29, 72) # Rose

            if item.get('como_corrigir'):
                p_corr = doc.add_paragraph()
                p_corr.paragraph_format.left_indent = Inches(0.2)
                r_co = p_corr.add_run(f"Como Ajustar: {item['como_corrigir']}")
                r_co.font.name = "Arial"
                r_co.font.size = Pt(9.5)
                r_co.font.color.rgb = RGBColor(16, 185, 129) # Emerald

    # Robôs ATS
    ats = data.get('analise_ats', {})
    if isinstance(ats, dict):
        doc.add_paragraph().paragraph_format.space_after = Pt(10)
        h_ats = doc.add_paragraph()
        r_hats = h_ats.add_run("COMPATIBILIDADE COM ROBÔS ATS (TRIAGEM AUTOMÁTICA)")
        r_hats.font.name = "Arial"
        r_hats.font.size = Pt(12)
        r_hats.font.bold = True
        r_hats.font.color.rgb = RGBColor(15, 23, 42)

        if ats.get('veredito_robos'):
            p_v = doc.add_paragraph()
            r_v = p_v.add_run(f"Veredito dos Robôs: {ats['veredito_robos']}")
            r_v.font.name = "Arial"
            r_v.font.size = Pt(10)

        if ats.get('palavras_chave_faltantes'):
            p_kw = doc.add_paragraph()
            r_kw = p_kw.add_run(f"Palavras-Chave Faltantes: {', '.join(ats['palavras_chave_faltantes'])}")
            r_kw.font.name = "Arial"
            r_kw.font.size = Pt(9.5)
            r_kw.font.bold = True
            r_kw.font.color.rgb = RGBColor(217, 119, 6)

    target = BytesIO()
    doc.save(target)
    target.seek(0)
    return target


def generate_pdf_export(data: dict, filename: str, job_target: str, model_name: str) -> BytesIO:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#0F172A')
    )

    subtitle_style = ParagraphStyle(
        'DocSubTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#6366F1')
    )

    section_heading = ParagraphStyle(
        'SecHeading',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=16,
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=12,
        spaceAfter=6
    )

    body_text = ParagraphStyle(
        'BodyDark',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor('#334155')
    )

    problem_text = ParagraphStyle(
        'ProbText',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#E11D48')
    )

    solution_text = ParagraphStyle(
        'SolText',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#059669')
    )

    elements = []

    # Header
    elements.append(Paragraph("TRIVOR", title_style))
    elements.append(Paragraph(f"TRIVOR  •  {filename}  •  {job_target}", subtitle_style))
    elements.append(Spacer(1, 14))

    # Score Card Banner Tabela
    score_val = f"{data.get('nota', 0.0):.1f}"
    score_html = f"<font size=28 color='#6366F1'><b>{score_val}</b></font><br/><font size=8 color='#94A3B8'><b>RESUME SCORE / 10</b></font>"
    parecer_html = f"<font size=9 color='#6366F1'><b>PARECER EXECUTIVO DO DIAGNÓSTICO</b></font><br/><br/><font size=9 color='#CBD5E1'>{data.get('resumo_executivo', '')}</font>"

    p_score = Paragraph(score_html, ParagraphStyle('ScP', alignment=1))
    p_parecer = Paragraph(parecer_html, ParagraphStyle('ParP'))

    t_data = [[p_score, p_parecer]]
    t_score = Table(t_data, colWidths=[130, 400])
    t_score.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (0,0), colors.HexColor('#0F172A')),
        ('BACKGROUND', (1,0), (1,0), colors.HexColor('#1E293B')),
        ('ALIGN', (0,0), (0,0), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('TOPPADDING', (0,0), (-1,-1), 12),
        ('BOTTOMPADDING', (0,0), (-1,-1), 12),
        ('LEFTPADDING', (0,0), (-1,-1), 12),
        ('RIGHTPADDING', (0,0), (-1,-1), 12),
    ]))

    elements.append(t_score)
    elements.append(Spacer(1, 14))

    # Diagnóstico por Seções
    elements.append(Paragraph("DIAGNÓSTICO ESTRUTURADO POR SEÇÃO", section_heading))
    secoes = data.get('diagnostico_por_secao', {})
    titulos = {
        'dados_pessoais': '👤 Dados Pessoais & Contatos',
        'resumo_profissional': '📝 Resumo / Perfil Profissional',
        'experiencia_profissional': '💼 Experiência Profissional & Métricas',
        'educacao_e_cursos': '🎓 Formação Acadêmica & Cursos',
        'habilidades_e_keywords': '🛠️ Habilidades & Palavras-Chave'
    }

    for k, title in titulos.items():
        if k in secoes and secoes[k]:
            item = secoes[k]
            status = item.get('status', 'atencao').upper()
            status_color = '#059669' if status == 'OK' else ('#D97706' if status == 'ATENCAO' else '#E11D48')

            sec_title = f"<b>{title}</b> &nbsp;&nbsp;<font size=8 color='{status_color}'><b>[{status}]</b></font>"
            elements.append(Paragraph(sec_title, body_text))

            if item.get('problema'):
                elements.append(Paragraph(f"• <b>Problema Detectado:</b> {item['problema']}", problem_text))
            if item.get('como_corrigir'):
                elements.append(Paragraph(f"• <b>Como Ajustar:</b> {item['como_corrigir']}", solution_text))

            elements.append(Spacer(1, 6))

    # ATS
    ats = data.get('analise_ats', {})
    if isinstance(ats, dict):
        elements.append(Spacer(1, 8))
        elements.append(Paragraph("COMPATIBILIDADE COM ROBÔS ATS", section_heading))
        if ats.get('veredito_robos'):
            elements.append(Paragraph(f"<b>Veredito dos Robôs:</b> {ats['veredito_robos']}", body_text))
        if ats.get('palavras_chave_faltantes'):
            elements.append(Paragraph(f"<b>Palavras-Chave Faltantes:</b> {', '.join(ats['palavras_chave_faltantes'])}", body_text))

    doc.build(elements)
    buffer.seek(0)
    return buffer
